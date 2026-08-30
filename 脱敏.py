#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件自动脱敏工具（MVP：正则 + 校验算法 + 实体清单，不上 NER）

用法：
  脱敏：  python3 脱敏.py <文件或目录> --mode ai|external [--entities 实体清单.txt] [--out 输出目录]
  反解：  python3 脱敏.py --restore 映射表.json <脱敏后文件>

模式：
  ai       —— 占位符替换（<PERSON_1> 等），同实体全文一致，生成映射表，可反解。用于给 AI 分析。
  external —— 掩码替换（张*、138****5678），不可逆，不生成映射表。用于对外提交。

支持格式：.docx / .pdf（文本层或扫描件，扫描件走 OCR）/ .txt / .md
输出：docx 输入 → 脱敏版 docx；PDF/txt/md 输入 → 脱敏版 txt
"""

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path

# ---------------- 识别：正则 + 校验算法 ----------------

def luhn_ok(digits: str) -> bool:
    """Luhn 算法，用于银行卡号校验"""
    total, alt = 0, False
    for ch in reversed(digits):
        d = ord(ch) - 48
        if alt:
            d *= 2
            if d > 9:
                d -= 9
        total += d
        alt = not alt
    return total % 10 == 0


def idcard_ok(s: str) -> bool:
    """18 位居民身份证校验位验证（GB 11643）[模型知识——需验证]"""
    weights = (7, 9, 10, 5, 8, 4, 2, 1, 6, 3, 7, 9, 10, 5, 8, 4, 2)
    check_map = "10X98765432"
    body = s[:17].upper()
    if not body.isdigit():
        return False
    checksum = sum(w * int(d) for w, d in zip(weights, body)) % 11
    return check_map[checksum] == s[-1].upper()


# 身份证形状（日期段合法的 18 位），用于银行卡规则排除
IDCARD_SHAPE = re.compile(
    r"^[1-9]\d{5}(?:19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx]$")


def bankcard_ok(s: str) -> bool:
    """Luhn 通过且不是身份证形状（日期形状的 18 位串按身份证处理）"""
    return luhn_ok(s) and not IDCARD_SHAPE.match(s)


# 各识别器：(类型, 优先级[小者优先], 正则, 校验函数或None)
DETECTORS = [
    ("IDCARD", 0,
     re.compile(r"(?<!\d)[1-9]\d{5}(?:19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx](?!\d)"),
     idcard_ok),
    ("BANKCARD", 1,
     re.compile(r"(?<!\d)\d{13,19}(?!\d)"),
     bankcard_ok),
    ("PHONE", 2,
     re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)"),
     None),
    ("LANDLINE", 3,
     re.compile(r"(?<!\d)0\d{2,3}-?\d{7,8}(?!\d)"),
     None),
    ("EMAIL", 4,
     re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"),
     None),
    ("USCC", 5,
     re.compile(r"(?<![0-9A-Za-z])[0-9]{2}[0-9A-HJ-NPQRTUWXY]{16}(?![0-9A-Za-z])"),
     None),
    ("IDCARD15", 6,  # 15 位老身份证，无校验位，低优先级
     re.compile(r"(?<!\d)[1-9]\d{5}\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}(?!\d)"),
     None),
    ("IDCARD_W", 7,  # 兜底：形状像身份证但校验位不过（证件号印错/OCR 误差），宁可多脱敏
     re.compile(r"(?<!\d)[1-9]\d{5}(?:19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx](?!\d)"),
     None),
    ("LONGNUM", 8,  # 兜底：15-19 位数字串（对公账号等未过 Luhn 的账号），宁可多脱敏
     re.compile(r"(?<!\d)\d{15,19}(?!\d)"),
     None),
]

# 疑似实体提示（仅输出到报告，不自动脱敏）——弥补不上 NER 的召回缺口
LABEL_HINT = re.compile(
    r"(原告|被告|上诉人|被上诉人|申请人|被申请人|甲方|乙方|丙方|丁方|戊方|己方"
    r"|委托人|受托人|发包人|承包人|出卖人|买受人|借款人|贷款人|出租人|承租人|保证人|债权人|债务人)"
    r"\s*[:：]?\s*([一-龥]{2,15})"
)
LAWYER_HINT = re.compile(r"([一-龥]{2,4})(?:律师)")
LONG_DIGIT_HINT = re.compile(r"(?<!\d)\d{10,}(?!\d)")


def find_regex_entities(text: str):
    """收集所有正则命中，重叠冲突时保留更长/更高优先级者"""
    cands = []
    for typ, prio, pat, checker in DETECTORS:
        for m in pat.finditer(text):
            s = m.group(0)
            if checker and not checker(s):
                continue
            cands.append((m.start(), m.end(), s, typ, prio))
    cands.sort(key=lambda c: (-(c[1] - c[0]), c[4]))
    chosen, occupied = [], set()
    for s, e, txt, typ, _ in cands:
        if any(i in occupied for i in range(s, e)):
            continue
        chosen.append((s, e, txt, typ))
        occupied.update(range(s, e))
    chosen.sort()
    return chosen


# ---------------- 识别：实体清单 ----------------

def guess_type(name: str) -> str:
    if re.search(r"(公司|企业|厂|集团|事务所|中心|银行|信用社|合作社|医院|学校|研究院)$", name):
        return "ORG"
    if re.search(r"(路|街|巷|弄|幢|号楼|号院|室|大厦|广场|小区|园区)", name):
        return "ADDR"
    return "PERSON"


def load_entities(path: Path):
    """返回 [(名称, 类型)]，长者优先匹配"""
    if not path.exists():
        return []
    items = []
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if "|" in line:
            name, _, typ = line.rpartition("|")
            items.append((name.strip(), typ.strip().upper()))
        else:
            items.append((line, guess_type(line)))
    items.sort(key=lambda x: -len(x[0]))
    return items


def find_list_entities(text: str, items):
    pats = []
    for name, typ in items:
        pats.append((re.escape(name), typ))
    if not pats:
        return []
    big = re.compile("|".join(f"(?P<g{i}>{p})" for i, (p, _) in enumerate(pats)))
    hits = []
    for m in big.finditer(text):
        for i, (p, typ) in enumerate(pats):
            if m.group(f"g{i}") is not None:
                hits.append((m.start(), m.end(), m.group(f"g{i}"), typ))
                break
    # 清单与正则结果重叠时，清单优先（调用方处理）
    return hits


# ---------------- 替换策略 ----------------

CN_NUM = "甲乙丙丁戊己庚辛壬癸"


def mask_external(typ: str, s: str, state: dict) -> str:
    """对外提交模式：掩码，不可逆"""
    if typ.startswith("IDCARD"):
        return s[:6] + "*" * (len(s) - 10) + s[-4:]
    if typ == "PHONE":
        return s[:3] + "****" + s[-4:]
    if typ == "LANDLINE":
        return s[:4] + "*" * (len(s) - 6) + s[-2:] if len(s) > 8 else s[:2] + "****"
    if typ == "BANKCARD" or typ == "LONGNUM":
        return "*" * (len(s) - 4) + s[-4:]
    if typ == "EMAIL":
        name, _, domain = s.partition("@")
        return (name[:1] + "***@" + domain) if name else "***@" + domain
    if typ == "USCC":
        return s[:6] + "*" * (len(s) - 10) + s[-4:]
    if typ == "PERSON":
        if len(s) <= 2:
            return s[0] + "*"
        return s[0] + "*" * (len(s) - 2) + s[-1]
    if typ == "ORG":
        n = state["org"]
        state["org"] += 1
        return f"某{CN_NUM[n]}公司" if n < 10 else f"某某公司{n}"
    if typ == "ADDR":
        return s[:6] + "****" if len(s) > 10 else s[:1] + "****"
    return s[0] + "***"


def desensitize_text(text: str, entity_items, mode: str, state=None):
    """返回 (脱敏后文本, 报告dict)。state 为文件级共享状态，保证同文件跨段落占位符/掩码一致"""
    if state is None:
        state = {"forward": {}, "reverse": {}, "counters": {}, "org": 0}
    # 1) 收集命中：实体清单优先于正则（人名可能是数字/字母以外任意文字，天然少冲突）
    hits = []
    taken = set()
    for s, e, val, typ in find_list_entities(text, entity_items):
        hits.append((s, e, val, typ))
        taken.update(range(s, e))
    for s, e, val, typ in find_regex_entities(text):
        if any(i in taken for i in range(s, e)):
            continue
        hits.append((s, e, val, typ))
    hits.sort()

    # 2) 替换
    report = {"命中数": len(hits)}
    if mode == "ai":
        def fn(val, typ):
            if val not in state["forward"]:
                state["counters"][typ] = state["counters"].get(typ, 0) + 1
                ph = f"<{typ}_{state['counters'][typ]}>"
                state["forward"][val] = ph
                state["reverse"][ph] = val
            return state["forward"][val]
    else:
        def fn(val, typ):
            return mask_external(typ, val, state)
    hits = [(s, e, v, t, fn) for s, e, v, t in hits]

    out = []
    last = 0
    for s, e, v, t in [(h[0], h[1], h[2], h[3]) for h in hits]:
        out.append(text[last:s])
        out.append(fn(v, t))
        last = e
    out.append(text[last:])
    new_text = "".join(out)

    # 3) 报告：疑似实体提示 + 未脱敏长数字串提示
    known = {h[2] for h in hits}
    field_words = re.compile(r"(账号|帐号|邮箱|信箱|电话|传真|地址|代表|代理人|号码|日期|联系人|负责人|经办人|编码)")
    hints = []
    for m in LABEL_HINT.finditer(text):
        val = m.group(2)
        if (val not in known and not field_words.search(val)
                and len(val) <= 6 and not re.search(r"\d", val)):
            hints.append(f"{m.group(1)}:{val}")
    for m in LAWYER_HINT.finditer(text):
        if m.group(1) not in known:
            hints.append(f"×律师:{m.group(1)}")
    leftovers = [m.group(0) for m in LONG_DIGIT_HINT.finditer(text)
                 if not any(h[0] <= m.start() < h[1] for h in hits)]
    report["疑似待确认实体"] = sorted(set(hints))
    report["未脱敏长数字串（请人工确认是否敏感）"] = sorted(set(leftovers))
    return new_text, report


CN_PUNCT = "，。；：、（）【】《》‘’“”！？"
ZH = r"一-龥"


def normalize_spaces(s: str) -> str:
    """去掉 PDF 抽取产生的无意义空格（汉字/中文标点两侧、数字字母与汉字之间），
    保留英文单词之间的空格。否则跨行实体与清单匹配不上"""
    cls = ZH + re.escape(CN_PUNCT)
    s = re.sub(rf"(?<=[{cls}])\s+(?=[{cls}])", "", s)
    s = re.sub(rf"(?<=[{cls}])\s+(?=[0-9A-Za-z])", "", s)
    s = re.sub(rf"(?<=[0-9A-Za-z])\s+(?=[{cls}])", "", s)
    return s


# ---------------- 文本抽取与输出 ----------------

def extract_docx(path: Path):
    """抽取 docx 全部文本：正文段落 + 表格单元格。返回段列表"""
    from docx import Document
    doc = Document(str(path))
    paras = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    paras.append(p.text)
    return paras


def extract_pdf(path: Path, ocr_dpi: int = 200):
    """抽取 PDF 文本；文本层过薄时走 OCR。
    返回 (段列表, 是否OCR)。每页合并为一个段落（判决书等文档实体常跨行断开，
    逐行处理会漏抓），并去掉页码行与无意义空格"""
    import fitz
    page_num_re = re.compile(r"^[-—–]\s*\d+\s*[-—–]$")  # - 2 - 型页码行

    def clean_page_lines(lines):
        kept = [ln.strip() for ln in lines
                if ln.strip() and not page_num_re.match(ln.strip())]
        return normalize_spaces("".join(kept))

    doc = fitz.open(str(path))
    paras = [clean_page_lines(page.get_text("text").splitlines()) for page in doc]
    doc.close()
    paras = [p for p in paras if p]
    if sum(len(p) for p in paras) >= 50:
        return paras, False
    # 文本层太薄 → 扫描件，OCR
    try:
        import numpy as np
        import cv2
        from rapidocr_onnxruntime import RapidOCR
    except ImportError as exc:
        raise SystemExit(
            f"「{path.name}」疑似扫描件，需要 OCR 依赖但未安装（{exc}）。\n"
            "请先执行：pip3 install rapidocr-onnxruntime"
        )
    ocr = RapidOCR()
    doc = fitz.open(str(path))
    paras = []
    for page in doc:
        pix = page.get_pixmap(dpi=ocr_dpi)
        img = cv2.imdecode(np.frombuffer(pix.tobytes("png"), np.uint8), cv2.IMREAD_COLOR)
        result, _ = ocr(img)
        if result:
            paras.append(normalize_spaces("".join(r[1] for r in result)))
    doc.close()
    return paras, True


def read_any(path: Path):
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        paras, _is_ocr = extract_pdf(path)
        return paras
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="replace").splitlines()
    raise SystemExit(f"不支持的格式：{path.name}（支持 .docx / .pdf / .txt / .md；老 .doc 请先另存为 .docx）")


def write_docx(paras, out_path: Path):
    from docx import Document
    doc = Document()
    for t in paras:
        doc.add_paragraph(t)
    doc.save(str(out_path))


def process_file(path: Path, entity_items, mode: str, out_dir: Path):
    paras = read_any(path)
    is_docx_in = path.suffix.lower() == ".docx"
    # 文件级共享状态：同文件内同一实体占位符/掩码一致（跨段落、跨表格）
    state = {"forward": {}, "reverse": {}, "counters": {}, "org": 0}
    new_paras, reports = [], []
    for line in paras:
        nt, rep = desensitize_text(line, entity_items, mode, state)
        new_paras.append(nt)
        reports.append(rep)

    stem = "脱敏_" + path.stem
    if is_docx_in:
        out_path = out_dir / (stem + ".docx")
        write_docx(new_paras, out_path)
    else:
        out_path = out_dir / (stem + ".txt")
        out_path.write_text("\n".join(new_paras), encoding="utf-8")

    # 映射表（仅 AI 模式）——保密级别与原文件相同
    map_path = None
    if mode == "ai":
        map_path = out_dir / (stem + ".映射表.json")
        map_path.write_text(json.dumps({
            "模式": "ai（占位符，可反解）",
            "生成时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "源文件": str(path),
            "正向": state["forward"],
            "反向": state["reverse"],
        }, ensure_ascii=False, indent=2), encoding="utf-8")

    # 汇总报告
    total = sum(r["命中数"] for r in reports)
    hint_set, digit_set = set(), set()
    for r in reports:
        hint_set.update(r["疑似待确认实体"])
        digit_set.update(r["未脱敏长数字串（请人工确认是否敏感）"])
    print(f"\n=== {path.name} ===")
    print(f"  输出：{out_path.name}（脱敏 {total} 处）")
    if map_path:
        print(f"  映射表：{map_path.name}（保密级别等同原文件，勿随脱敏件外发）")
    if hint_set:
        print("  疑似待确认实体（正则抓不到，建议确认后加入实体清单）：")
        for h in sorted(hint_set):
            print(f"    - {h}")
    if digit_set:
        print("  未脱敏长数字串（可能是账号/编号，请人工确认）：")
        for d in sorted(digit_set):
            print(f"    - {d}")
    return out_path, map_path


def restore(map_file: Path, masked_file: Path):
    """用映射表把占位符还原为原文"""
    data = json.loads(map_file.read_text(encoding="utf-8"))
    reverse = data["反向"]
    text = read_any(masked_file)
    text = [reverse.get(line, line) for line in text]
    # 占位符可能嵌在长行中，再做一遍整体替换
    def sub_line(s):
        for ph, val in reverse.items():
            s = s.replace(ph, val)
        return s
    text = [sub_line(line) for line in text]
    out = masked_file.parent / ("还原_" + masked_file.stem.replace("脱敏_", "") + masked_file.suffix)
    if masked_file.suffix == ".docx":
        write_docx(text, out)
    else:
        out.write_text("\n".join(text), encoding="utf-8")
    print(f"已还原 → {out}")


def main():
    ap = argparse.ArgumentParser(description="文件自动脱敏（法律文书场景 MVP）")
    ap.add_argument("input", nargs="?", help="输入文件或目录")
    ap.add_argument("--mode", choices=["ai", "external"], default="ai",
                    help="ai=占位符可逆（给AI分析）；external=掩码不可逆（对外提交）")
    ap.add_argument("--entities", default=str(Path(__file__).parent / "实体清单.txt"),
                    help="实体清单文件（每行一个，# 为注释）")
    ap.add_argument("--out", default=None, help="输出目录，默认为输入所在目录")
    ap.add_argument("--restore", metavar="映射表.json", help="反解模式：指定映射表")
    args = ap.parse_args()

    if args.restore:
        if not args.input:
            ap.error("反解模式需同时给出脱敏后文件路径")
        restore(Path(args.restore), Path(args.input))
        return
    if not args.input:
        ap.error("请给出输入文件或目录（或用 --restore 反解）")

    inp = Path(args.input)
    files = [inp] if inp.is_file() else sorted(
        p for p in inp.iterdir() if p.suffix.lower() in (".docx", ".pdf", ".txt", ".md")
        and not p.name.startswith("脱敏_"))
    if not files:
        ap.error("未找到可处理的文件")
    out_dir = Path(args.out) if args.out else files[0].parent
    out_dir.mkdir(parents=True, exist_ok=True)

    entity_items = load_entities(Path(args.entities))
    print(f"模式：{'AI分析（占位符，可反解）' if args.mode == 'ai' else '对外提交（掩码，不可逆）'}"
          f"｜实体清单：{len(entity_items)} 项")
    for f in files:
        try:
            process_file(f, entity_items, args.mode, out_dir)
        except SystemExit as exc:
            # 单个文件失败（如缺 OCR 依赖）不中止整批
            print(f"[跳过] {f.name}：{exc}")


if __name__ == "__main__":
    main()
